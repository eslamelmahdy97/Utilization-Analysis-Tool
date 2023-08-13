import pandas as pd 
import numpy as np
import datetime as dt
import docx
from docx import Document
from docx.shared import Inches
from docx.shared import Cm

selectmon=input('Enter The Month of choice : ')
selectmon = int(selectmon)
accountname=input('Please Enter The Account name:')
csvfile=accountname+'.csv'
docfile=accountname+'.docx'

# create a new word document
document = docx.Document()

# add heading
document.add_heading(accountname+" Insights", 0)

df=pd.read_csv(csvfile)

#Data Preparation
df['AMOUNT PAID BY RISK CARRIER']=pd.to_numeric(df['AMOUNT PAID BY RISK CARRIER'])
df['CLAIM DATE'] = pd.to_datetime(df['CLAIM DATE'], format='%d-%m-%y')
df['DATE OF BIRTH']=pd.to_datetime(df['DATE OF BIRTH'])
df['claimmonth'] = pd.DatetimeIndex(df['CLAIM DATE']).month
df['claimmonth'] = pd.to_datetime(df['CLAIM DATE']).dt.month

# The whole analysis paragraph
def analysis_paramaters(df):
    #number of (claims)
    rcrdcount=len(df)
    #total Consumption
    totalconsumption = round(sum(df['AMOUNT PAID BY RISK CARRIER']))
    formatted_total = "{:,}".format(totalconsumption)
    # Average Claims Severity
    avgclaimsev=round(totalconsumption/len(df[df['AMOUNT PAID BY RISK CARRIER']!=0]))
    # Distinct count of members
    membercount=len(pd.unique(df['INDIVIDUAL#']))
    #average Claims Frequency
    try:
        avgclaimfreq=round(rcrdcount/membercount)
    except:
        avgclaimfreq="no entries"
    # Diagnosis Breakdown
    topdiag= df.groupby('INITIAL/MAJOR DISEASE')['AMOUNT PAID BY RISK CARRIER'].agg(['sum','count'])
    topdiag = topdiag.sort_values(by = ['sum'], ascending=[False]).head(2)
    # create a formatted version of the topdiag dataframe
    #formatted_topdiag = topdiag.applymap(lambda x: f'{x:,.0f}' if isinstance(x, (int, float)) else x)

# print the formatted dataframe
   
    #Providers Breakdown
    topprov= df.groupby('PROVIDER')['AMOUNT PAID BY RISK CARRIER'].agg(['sum','count'])
    topprov = topprov.sort_values(by = ['sum'], ascending=[False]).head(2)
    #Top Months Breakdown
    #print('Highest Months: \n {} '.format(topmon), file=text_file)
    
    # add analysis parameters to the document
    p = document.add_paragraph()
    p.add_run('The total consumption is ')
    run = p.add_run('{} EGP'.format(formatted_total))
    run.bold = True
    p.add_run(' with total Claims count of ')
    run = p.add_run('{} Claims.\n'.format(rcrdcount))
    run.bold = True
    p.add_run('Average claim frequency is ')
    run = p.add_run('{} claims'.format(avgclaimfreq))
    run.bold = True
    p.add_run(', While Average claim severity is ')
    run = p.add_run('{} EGP.\n'.format(avgclaimsev))
    run.bold = True
    p.add_run('Total number of using members is ')
    run = p.add_run('{} members.'.format(membercount))
    run.bold = True
    
    if not topdiag.empty:
    
        table = document.add_table(rows=1, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Disease'
        hdr_cells[1].text = 'Total Cost'
        hdr_cells[2].text = 'Count of Claims'
        hdr_cells[0].width = Inches(4)
        for idx, row in topdiag.iterrows():
            row_cells = table.add_row().cells
            row_cells[0].text = idx
            row_cells[1].text = '{:,.0f}'.format(row['sum'])
            row_cells[2].text = '{:,.0f}'.format(row['count'])
            row_cells[0].width = Inches(4)
        table.style = 'Table Grid'
    
    if not topprov.empty:
        document.add_paragraph()
        table = document.add_table(rows=1, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Provider'
        hdr_cells[1].text = 'Total Cost'
        hdr_cells[2].text = 'Count of Claims'
        hdr_cells[0].width = Inches(4)
        for idx, row in topprov.iterrows():
            row_cells = table.add_row().cells
            row_cells[0].text = idx
            row_cells[1].text = '{:,.0f}'.format(row['sum'])
            row_cells[2].text = '{:,.0f}'.format(row['count'])
            row_cells[0].width = Inches(4)
        table.style = 'Table Grid'
    
# Cumulative Consumption
analysis_paramaters(df)

# a list of all subcategories sorted by claims count    
tsub = df['SUB-CATEGORY'].value_counts().index.tolist()
tsub_sorted = sorted(tsub, key=lambda x: df[df['SUB-CATEGORY']==x]['AMOUNT PAID BY RISK CARRIER'].sum(), reverse=True)

#subcategories Breakdown 
for sub in tsub_sorted:
    document.add_heading(sub + " Insights", level=1)
    
    dftemp=df[df['SUB-CATEGORY']==sub]
    analysis_paramaters(dftemp)
    if sub == 'Prescription Medicine':
        document.add_paragraph("Chronic Insights")
        
        dftemp1=dftemp[df['Chronic_Posting']=='YES']
        analysis_paramaters(dftemp1)
        
        dftemp2=dftemp[df['Chronic_Posting']=='NO']
        document.add_paragraph("Acute Insights")
        analysis_paramaters(dftemp2)
        
    
# Last Months Consumption
dftemp3 = df[df['CLAIM DATE'].dt.month == selectmon]
document.add_heading("Last Month's Insights", level=1)
analysis_paramaters(dftemp3)
section = document.sections[0]
section.left_margin = Inches(0.5)
section.right_margin = Inches(0.5)
section.top_margin = Inches(0.5)
section.bottom_margin = Inches(0.5)
header = section.header
paragraph = header.add_paragraph()
paragraph.alignment = 2 # right-align the paragraph
run = paragraph.add_run()
run.add_picture('logo.png', width=Cm(4.2), height=Cm(1.75))
document.save(docfile)
